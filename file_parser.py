import io

import pdfplumber
from docx import Document


def _join_non_empty(parts: list[str]) -> str:
    return "\n".join(part for part in parts if str(part or "").strip()).strip()


def _dedupe_preserve_order(lines: list[str]) -> list[str]:
    out: list[str] = []
    seen: set[str] = set()
    for line in lines:
        key = str(line or "").strip()
        if not key or key in seen:
            continue
        seen.add(key)
        out.append(key)
    return out


def _extract_pdf_page_words(page: pdfplumber.page.Page) -> str:
    """
    Word-level fallback that captures text missed by simple extract_text(),
    especially in multi-column layouts.
    """
    words = page.extract_words(use_text_flow=False, keep_blank_chars=False) or []
    if not words:
        return ""

    y_bucket = 3.0
    buckets: dict[int, list[dict]] = {}
    for word in words:
        top = float(word.get("top", 0.0))
        key = int(round(top / y_bucket))
        buckets.setdefault(key, []).append(word)

    lines: list[str] = []
    for key in sorted(buckets.keys()):
        row = sorted(buckets[key], key=lambda w: float(w.get("x0", 0.0)))
        text = " ".join(str(item.get("text") or "").strip() for item in row).strip()
        if text:
            lines.append(text)
    return "\n".join(lines).strip()


def parse_pdf(file_bytes: bytes) -> str:
    pages: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            primary = (page.extract_text(layout=True) or "").strip()
            fallback = _extract_pdf_page_words(page)
            if primary and fallback and primary != fallback:
                pages.append(_join_non_empty([primary, fallback]))
            else:
                pages.append(primary or fallback)
    return _join_non_empty(pages)


def parse_docx(file_bytes: bytes) -> str:
    document = Document(io.BytesIO(file_bytes))
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = str(paragraph.text or "").strip()
        if text:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = str(paragraph.text or "").strip()
                    if text:
                        lines.append(text)

    return _join_non_empty(_dedupe_preserve_order(lines))


def parse_txt(file_bytes: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1252"):
        try:
            return file_bytes.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="ignore").strip()


def extract_text(file_bytes: bytes, filename: str) -> str:
    extension = filename.lower()

    if extension.endswith(".pdf"):
        return parse_pdf(file_bytes)
    if extension.endswith(".docx"):
        return parse_docx(file_bytes)
    if extension.endswith(".txt"):
        return parse_txt(file_bytes)

    raise ValueError("Unsupported file format")
