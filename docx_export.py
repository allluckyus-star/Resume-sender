import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SECTION_RENDER_ORDER = ("summary", "experience", "skills", "education", "projects")
PAGE_WIDTH = Inches(8.5)
PAGE_HEIGHT = Inches(11)
SIDE_MARGIN = Inches(0.75)
TOP_BOTTOM_MARGIN = Inches(0.75)
NAME_SIZE = Pt(18)
HEADLINE_SIZE = Pt(13)
HEADING_SIZE = Pt(15)
ROLE_SIZE = Pt(13)
BODY_SIZE = Pt(11)
SKILL_SIZE = Pt(11)
CONTACT_SIZE = Pt(10.5)
META_SIZE = Pt(10.5)
STANDARD_FONT = "Calibri"
LINE_SPACING = 1.1
MAX_HEADLINE_CHARS = 86
MAX_CONTACT_CHARS = 96
DATE_SEPARATOR_PATTERN = re.compile(r"\s*[-–—]\s*")
PRIMARY_TEXT_COLOR = RGBColor(0x00, 0x00, 0x00)
SECTION_TEXT_COLOR = RGBColor(0x1F, 0x1F, 0x1F)
SECTION_TOP_SPACE = 14
SECTION_BOTTOM_SPACE = 8
ROLE_TOP_SPACE = 10
ROLE_BOTTOM_SPACE = 3
BULLET_BOTTOM_SPACE = 3
BULLET_LEFT_INDENT = Inches(0.3)

MONTH_LOOKUP = {
    "jan": "01",
    "january": "01",
    "feb": "02",
    "february": "02",
    "mar": "03",
    "march": "03",
    "apr": "04",
    "april": "04",
    "may": "05",
    "jun": "06",
    "june": "06",
    "jul": "07",
    "july": "07",
    "aug": "08",
    "august": "08",
    "sep": "09",
    "sept": "09",
    "september": "09",
    "oct": "10",
    "october": "10",
    "nov": "11",
    "november": "11",
    "dec": "12",
    "december": "12",
}
MONTH_NUMBER_TO_ABBR = {
    "01": "Jan",
    "02": "Feb",
    "03": "Mar",
    "04": "Apr",
    "05": "May",
    "06": "Jun",
    "07": "Jul",
    "08": "Aug",
    "09": "Sep",
    "10": "Oct",
    "11": "Nov",
    "12": "Dec",
}


def _add_spacer(document: Document, *, size: Pt, after: float = 0) -> None:
    paragraph = document.add_paragraph()
    _style_paragraph(paragraph, after=after, line_spacing=LINE_SPACING)
    run = paragraph.add_run("")
    run.font.size = size


def _normalize_url(url: str) -> str:
    value = str(url).strip()
    if not value:
        return ""
    if value.startswith(("http://", "https://", "mailto:", "tel:")):
        return value
    if "@" in value and " " not in value:
        return f"mailto:{value}"
    return f"https://{value}"


def _link_label(url: str) -> str:
    value = str(url).strip()
    lowered = value.lower()
    if "linkedin.com" in lowered:
        return "LinkedIn"
    if "github.com" in lowered:
        return "GitHub"
    if "portfolio" in lowered or any(domain in lowered for domain in (".dev", ".io", ".ai")):
        return "Portfolio"
    return "Website"


def _compact_text(value: str, max_chars: int) -> str:
    text = re.sub(r"\s+", " ", str(value or "")).strip(" |")
    if len(text) <= max_chars:
        return text

    parts = [part.strip() for part in text.split("|") if part.strip()]
    if len(parts) > 1:
        kept: list[str] = []
        for part in parts:
            candidate = " | ".join(kept + [part])
            if len(candidate) <= max_chars:
                kept.append(part)
        if kept:
            return " | ".join(kept)

    shortened = text[: max_chars + 1].rsplit(" ", 1)[0].strip(" ,|")
    return shortened or text[:max_chars].strip(" ,|")


def _add_hyperlink(paragraph, text: str, url: str, *, font_size: Pt) -> None:
    relation_id = paragraph.part.relate_to(_normalize_url(url), RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    run_properties.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)

    size = OxmlElement("w:sz")
    size.set(qn("w:val"), str(int(font_size.pt * 2)))
    run_properties.append(size)

    run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _format_single_date(value: str) -> str:
    text = str(value).strip()
    if not text:
        return ""
    if text.lower() == "present":
        return "Present"

    numeric = re.fullmatch(r"(0?[1-9]|1[0-2])/(\d{2,4})", text)
    if numeric:
        month, year = numeric.groups()
        full_year = year if len(year) == 4 else f"20{year}"
        month_value = f"{int(month):02d}"
        return f"{MONTH_NUMBER_TO_ABBR.get(month_value, month_value)} {full_year}"

    month_name = re.fullmatch(r"([A-Za-z]+)\s+(\d{4})", text)
    if month_name:
        month_word, year = month_name.groups()
        month = MONTH_LOOKUP.get(month_word.lower())
        if month:
            return f"{MONTH_NUMBER_TO_ABBR.get(month, month)} {year}"

    if re.fullmatch(r"\d{4}", text):
        return text
    return text


def _format_duration(value: str) -> str:
    text = str(value).strip()
    if not text:
        return ""
    parts = DATE_SEPARATOR_PATTERN.split(text)
    if len(parts) == 1:
        return _format_single_date(parts[0])
    if len(parts) >= 2:
        return f"{_format_single_date(parts[0])} - {_format_single_date(parts[1])}"
    return text


def _configure_page(document: Document) -> None:
    section = document.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = TOP_BOTTOM_MARGIN
    section.bottom_margin = TOP_BOTTOM_MARGIN
    section.left_margin = SIDE_MARGIN
    section.right_margin = SIDE_MARGIN


def _set_global_font(document: Document) -> None:
    style = document.styles["Normal"]
    font = style.font
    font.name = STANDARD_FONT
    font.size = BODY_SIZE
    font.color.rgb = PRIMARY_TEXT_COLOR

    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = LINE_SPACING
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)


def _set_run_font(run, *, size: Pt | None = None, italic: bool | None = None, bold: bool | None = None) -> None:
    run.font.name = STANDARD_FONT
    if size is not None:
        run.font.size = size
    if italic is not None:
        run.italic = italic
    if bold is not None:
        run.bold = bold


def _style_paragraph(paragraph, *, before: float = 0, after: float = 0, line_spacing: float = LINE_SPACING) -> None:
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(before)
    paragraph_format.space_after = Pt(after)
    paragraph_format.line_spacing = line_spacing


def _section_sort_key(section: dict) -> int:
    section_type = str(section.get("type") or "").strip().lower()
    try:
        return SECTION_RENDER_ORDER.index(section_type)
    except ValueError:
        return len(SECTION_RENDER_ORDER)


def _section_map(resume: dict) -> dict[str, dict]:
    """Map section type → section. Types are lowercased so \"Education\" from the model still renders."""
    out: dict[str, dict] = {}
    for section in resume.get("sections", []):
        if not isinstance(section, dict):
            continue
        key = str(section.get("type") or "").strip().lower()
        if not key:
            continue
        out[key] = section
    return out


def _add_name(document: Document, name: str) -> None:
    if not name:
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _style_paragraph(paragraph, after=2, line_spacing=LINE_SPACING)
    run = paragraph.add_run(name)
    run.bold = True
    run.font.size = NAME_SIZE
    run.font.color.rgb = PRIMARY_TEXT_COLOR


def _add_headline(document: Document, headline: str) -> None:
    if not headline:
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _style_paragraph(paragraph, after=2, line_spacing=LINE_SPACING)
    run = paragraph.add_run(_compact_text(headline, MAX_HEADLINE_CHARS))
    run.font.size = HEADLINE_SIZE
    run.font.color.rgb = PRIMARY_TEXT_COLOR


def _add_second_headline(document: Document, headline2: str) -> None:
    return


def _add_contact_row(document: Document, header: dict) -> None:
    phone = str(header.get("phone", "")).strip()
    email = str(header.get("email", "")).strip()
    links = [part.strip() for part in str(header.get("links", "")).split("|") if part.strip()]
    location = str(header.get("location", "")).strip()
    fallback_contact = [part.strip() for part in str(header.get("contact", "")).split("|") if part.strip()]

    if not any([phone, email, links, location, fallback_contact]):
        return

    parts: list[tuple[str, str]] = []
    if email:
        parts.append(("email", email))
    for item in fallback_contact:
        if item and item not in {phone, email}:
            parts.append(("email" if "@" in item and " " not in item else "text", item))
    for link in links:
        parts.append(("link", link))
    if phone:
        parts.append(("text", phone))
    if location:
        parts.append(("text", location))

    def display_text(part: tuple[str, str]) -> str:
        kind, value = part
        return _link_label(value) if kind == "link" else value

    while len(" | ".join(display_text(part) for part in parts)) > MAX_CONTACT_CHARS and len(parts) > 1:
        priorities = {"link": 0, "email": 1, "text": 2}
        removable_index = min(range(len(parts)), key=lambda index: priorities.get(parts[index][0], 0))
        parts.pop(removable_index)

    for _, value in parts:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _style_paragraph(paragraph, after=1, line_spacing=LINE_SPACING)
        run = paragraph.add_run(value)
        run.font.size = CONTACT_SIZE
        run.font.color.rgb = PRIMARY_TEXT_COLOR


def _header_contact_parts(header: dict) -> list[tuple[str, str]]:
    phone = str(header.get("phone", "")).strip()
    email = str(header.get("email", "")).strip()
    links = [part.strip() for part in str(header.get("links", "")).split("|") if part.strip()]
    location = str(header.get("location", "")).strip()
    fallback_contact = [part.strip() for part in str(header.get("contact", "")).split("|") if part.strip()]

    # Strict order for two-column header rendering.
    parts: list[tuple[str, str]] = []
    if email:
        parts.append(("email", email))
    if links:
        parts.append(("link", links[0]))
    if phone:
        parts.append(("text", phone))
    if location:
        parts.append(("text", location))

    # Append remaining values in stable order after strict core fields.
    for link in links[1:]:
        if link:
            parts.append(("link", link))
    for item in fallback_contact:
        if item and item not in {phone, email, location} and item not in links:
            parts.append(("email" if "@" in item and " " not in item else "text", item))
    return parts


def _add_two_column_header(document: Document, header: dict) -> None:
    name = str(header.get("name", "")).strip()
    headline = str(header.get("headline", "")).strip()
    contacts = [value for _, value in _header_contact_parts(header)]
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(4.55)
    table.columns[1].width = Inches(2.45)

    left_cell = table.rows[0].cells[0]
    right_cell = table.rows[0].cells[1]

    for cell in (left_cell, right_cell):
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            element = OxmlElement(f"w:{edge}")
            element.set(qn("w:val"), "nil")
            tc_borders.append(element)
        tc_pr.append(tc_borders)
        cell.text = ""

    if name:
        paragraph = left_cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _style_paragraph(paragraph, after=2, line_spacing=LINE_SPACING)
        run = paragraph.add_run(name)
        _set_run_font(run, size=NAME_SIZE, bold=True)
        run.font.color.rgb = PRIMARY_TEXT_COLOR

    if headline:
        paragraph = left_cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _style_paragraph(paragraph, after=1, line_spacing=LINE_SPACING)
        run = paragraph.add_run(_compact_text(headline, MAX_HEADLINE_CHARS))
        _set_run_font(run, size=HEADLINE_SIZE)
        run.font.color.rgb = PRIMARY_TEXT_COLOR

    if contacts:
        # Add one empty line before the contact block (not between contact rows).
        right_cell.add_paragraph()

    for contact in contacts:
        paragraph = right_cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _style_paragraph(paragraph, after=1, line_spacing=LINE_SPACING)
        run = paragraph.add_run(contact)
        _set_run_font(run, size=CONTACT_SIZE, italic=True)
        run.font.color.rgb = PRIMARY_TEXT_COLOR


def _add_header_gap(document: Document) -> None:
    paragraph = document.add_paragraph()
    _style_paragraph(paragraph, after=12, line_spacing=LINE_SPACING)


def _add_section_heading(document: Document, title: str, *, before: float = 0, after: float = 6) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = "Heading 2"
    _style_paragraph(
        paragraph,
        before=before if before else SECTION_TOP_SPACE,
        after=max(after, SECTION_BOTTOM_SPACE),
        line_spacing=LINE_SPACING,
    )
    run = paragraph.add_run(str(title).strip())
    _set_run_font(run, size=HEADING_SIZE, bold=True)
    run.font.color.rgb = SECTION_TEXT_COLOR


def _add_subsection_divider(paragraph) -> None:
    return


def _indent_subsection_paragraph(paragraph) -> None:
    paragraph.paragraph_format.left_indent = Inches(0.0)


def _add_summary(document: Document, section: dict | None) -> None:
    if not section:
        return
    items = [
        str(item.get("text", "")).strip()
        for item in section.get("items", [])
        if isinstance(item, dict) and item.get("text")
    ]
    if not items:
        return
    paragraph = document.add_paragraph()
    _style_paragraph(paragraph, before=8, after=10)
    _add_text_with_bold_markers(paragraph, " ".join(items), BODY_SIZE)


def _add_role_duration_line(document: Document, role: str, duration: str, *, before: float = 10):
    if not role and not duration:
        return None
    paragraph = document.add_paragraph()
    _indent_subsection_paragraph(paragraph)
    _style_paragraph(paragraph, before=max(before, ROLE_TOP_SPACE), after=ROLE_BOTTOM_SPACE, line_spacing=LINE_SPACING)
    if role:
        role_run = paragraph.add_run(role)
        _set_run_font(role_run, size=ROLE_SIZE, bold=True)
        role_run.font.color.rgb = PRIMARY_TEXT_COLOR
    if duration:
        duration_run = paragraph.add_run(f"    {_format_duration(duration)}")
        _set_run_font(duration_run, size=ROLE_SIZE)
        duration_run.font.color.rgb = PRIMARY_TEXT_COLOR
    return paragraph


def _add_skill_line(document: Document, category: str, values: list[str]) -> None:
    if not category and not values:
        return
    paragraph = document.add_paragraph()
    _style_paragraph(paragraph, after=3, line_spacing=LINE_SPACING)

    if category:
        label = paragraph.add_run(f"{category}: ")
        _set_run_font(label, size=SKILL_SIZE, bold=True)

    if values:
        values_run = paragraph.add_run(", ".join(values))
        _set_run_font(values_run, size=SKILL_SIZE)


def _render_skills(document: Document, section: dict | None) -> None:
    if not section or not section.get("items"):
        return
    _add_section_heading(document, section.get("title", "Skills"), before=14)
    for item in section.get("items", []):
        if not isinstance(item, dict):
            continue
        category = str(item.get("category", "")).strip()
        values = [str(value).strip() for value in item.get("values", []) if str(value).strip()]
        if values or category:
            _add_skill_line(document, category, values)


def _render_key_achievements(document: Document, section: dict | None) -> None:
    if not section or not section.get("items"):
        return
    texts = [
        str(item.get("text", "") if isinstance(item, dict) else item).strip()
        for item in section.get("items", [])
    ]
    texts = [text for text in texts if text]
    if not texts:
        return
    _add_section_heading(document, section.get("title", "Key Achievements"), before=14)
    for text in texts:
        _add_bullet(document, text)


def _add_role_line(document: Document, text: str, *, before: float = 10) -> None:
    if not text:
        return None
    paragraph = document.add_paragraph()
    _indent_subsection_paragraph(paragraph)
    _style_paragraph(paragraph, before=before, after=2, line_spacing=LINE_SPACING)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = ROLE_SIZE
    run.font.color.rgb = PRIMARY_TEXT_COLOR
    return paragraph


def _add_meta_line(document: Document, text: str) -> None:
    if not text:
        return None
    paragraph = document.add_paragraph()
    _indent_subsection_paragraph(paragraph)
    _style_paragraph(paragraph, before=1, after=3, line_spacing=LINE_SPACING)
    run = paragraph.add_run(text)
    _set_run_font(run, size=META_SIZE, italic=True)
    run.font.color.rgb = PRIMARY_TEXT_COLOR
    return paragraph


def _add_education_meta_line(document: Document, school: str, year: str) -> None:
    if not school and not year:
        return None
    paragraph = document.add_paragraph()
    _indent_subsection_paragraph(paragraph)
    _style_paragraph(paragraph, before=2, after=2, line_spacing=LINE_SPACING)
    if school:
        school_run = paragraph.add_run(school)
        school_run.italic = True
        school_run.font.size = META_SIZE
    if school and year:
        separator = paragraph.add_run(" | ")
        separator.font.size = META_SIZE
    if year:
        year_run = paragraph.add_run(_format_duration(year))
        year_run.font.size = META_SIZE
    return paragraph


def _add_focus_line(document: Document, focus_area: str, focus_details: str) -> None:
    if not focus_area and not focus_details:
        return None
    paragraph = document.add_paragraph()
    _indent_subsection_paragraph(paragraph)
    _style_paragraph(paragraph, before=1, after=3, line_spacing=LINE_SPACING)
    if focus_area:
        label = paragraph.add_run(f"{focus_area}: ")
        label.bold = True
        label.font.size = BODY_SIZE
    if focus_details:
        details = paragraph.add_run(focus_details)
        details.font.size = BODY_SIZE
    return paragraph


def _add_duration_line(document: Document, duration: str) -> None:
    if not duration:
        return None
    paragraph = document.add_paragraph()
    _indent_subsection_paragraph(paragraph)
    _style_paragraph(paragraph, after=4, line_spacing=LINE_SPACING)
    run = paragraph.add_run(_format_duration(duration))
    run.font.size = META_SIZE
    return paragraph


def _add_overview_line(document: Document, overview: str) -> None:
    if not overview:
        return None
    paragraph = document.add_paragraph()
    _indent_subsection_paragraph(paragraph)
    _style_paragraph(paragraph, after=6, line_spacing=LINE_SPACING)
    run = paragraph.add_run(overview)
    run.italic = True
    return paragraph


def _add_text_with_bold_markers(paragraph, text: str, font_size) -> None:
    for part in re.split(r"(\*\*[^*]+?\*\*)", str(text or "")):
        if not part:
            continue
        is_bold = part.startswith("**") and part.endswith("**")
        value = part[2:-2] if is_bold else part
        if not value:
            continue
        run = paragraph.add_run(value)
        _set_run_font(run, size=font_size, bold=is_bold)


def _normalize_highlights(highlights) -> list[str]:
    if not isinstance(highlights, list):
        return []
    cleaned: list[str] = []
    for item in highlights:
        value = str(item or "").strip()
        if value and value not in cleaned:
            cleaned.append(value)
        if len(cleaned) >= 2:
            break
    return cleaned


def _add_text_with_highlights(paragraph, text: str, highlights: list[str], font_size) -> None:
    source = str(text or "")
    if not source:
        return

    spans: list[tuple[int, int]] = []
    for highlight in _normalize_highlights(highlights):
        start = source.find(highlight)
        if start >= 0:
            spans.append((start, start + len(highlight)))

    if not spans:
        _add_text_with_bold_markers(paragraph, source, font_size)
        return

    spans.sort(key=lambda item: item[0])
    merged: list[tuple[int, int]] = []
    for start, end in spans:
        if not merged or start >= merged[-1][1]:
            merged.append((start, end))
        else:
            merged[-1] = (merged[-1][0], max(merged[-1][1], end))

    cursor = 0
    for start, end in merged:
        if start > cursor:
            run = paragraph.add_run(source[cursor:start])
            _set_run_font(run, size=font_size)
        run = paragraph.add_run(source[start:end])
        _set_run_font(run, size=font_size, bold=True)
        cursor = end

    if cursor < len(source):
        run = paragraph.add_run(source[cursor:])
        _set_run_font(run, size=font_size)


def _add_bullet(document: Document, text):
    paragraph = document.add_paragraph()
    paragraph.style = "List Bullet"
    _style_paragraph(paragraph, after=BULLET_BOTTOM_SPACE, line_spacing=LINE_SPACING)
    # Hanging indent so wrapped lines align under bullet text, not under the dot.
    paragraph.paragraph_format.left_indent = Inches(0.35)
    paragraph.paragraph_format.first_line_indent = Inches(-0.15)

    if isinstance(text, dict):
        cleaned = str(text.get("text", "")).strip()
        highlights = _normalize_highlights(text.get("highlights", []))
        _add_text_with_highlights(paragraph, cleaned, highlights, BODY_SIZE)
    else:
        cleaned = str(text).strip()
        _add_text_with_bold_markers(paragraph, cleaned, BODY_SIZE)
    return paragraph


def _render_experience(document: Document, section: dict | None) -> None:
    if not section or not section.get("items"):
        return
    items = section.get("items", [])
    _add_section_heading(document, section.get("title", "Experience"), before=14)
    for index, item in enumerate(items):
        role = str(item.get("role", "")).strip()
        company = str(item.get("company", "")).strip()
        location = str(item.get("location", "")).strip()
        duration = str(item.get("duration", "")).strip()

        last_paragraph = _add_role_duration_line(document, role, duration)
        last_paragraph = _add_meta_line(document, " | ".join(part for part in (company, location) if part)) or last_paragraph
        for bullet in item.get("bullets", []):
            last_paragraph = _add_bullet(document, bullet)
        if index < len(items) - 1 and last_paragraph is not None:
            _add_subsection_divider(last_paragraph)


def _render_projects(document: Document, section: dict | None) -> None:
    if not section or not section.get("items"):
        return
    items = section.get("items", [])
    _add_section_heading(document, section.get("title", "Projects"), before=14)
    for index, item in enumerate(items):
        last_paragraph = _add_role_line(document, item.get("name", ""))
        last_paragraph = _add_meta_line(
            document,
            " | ".join(
                part for part in (", ".join(item.get("stack", [])), _format_duration(item.get("duration", ""))) if part
            ),
        ) or last_paragraph
        last_paragraph = _add_overview_line(document, item.get("overview", "")) or last_paragraph
        for bullet in item.get("bullets", []):
            last_paragraph = _add_bullet(document, bullet)
        if index < len(items) - 1 and last_paragraph is not None:
            _add_subsection_divider(last_paragraph)


def _render_education(document: Document, section: dict | None) -> None:
    if not section or not section.get("items"):
        return
    items = section.get("items", [])
    _add_section_heading(document, section.get("title", "Education"), before=14)
    for index, item in enumerate(items):
        degree = str(item.get("degree", "")).strip()
        field = str(item.get("field", "")).strip()
        school = str(item.get("school", "")).strip()
        year = str(item.get("duration", "")).strip()
        grade = str(item.get("grade", "") or item.get("gpa", "")).strip()

        line1 = document.add_paragraph()
        _style_paragraph(line1, before=8 if index > 0 else 0, after=1, line_spacing=LINE_SPACING)
        if school:
            school_run = line1.add_run(school)
            school_run.bold = True
            school_run.font.size = BODY_SIZE
        if year:
            spacer = "    " if school else ""
            year_run = line1.add_run(f"{spacer}{_format_duration(year)}")
            year_run.font.size = BODY_SIZE

        line2 = document.add_paragraph()
        _style_paragraph(line2, after=1, line_spacing=LINE_SPACING)
        if degree:
            degree_run = line2.add_run(degree)
            _set_run_font(degree_run, size=BODY_SIZE, bold=True)
        if field:
            prefix = " " if degree else ""
            field_run = line2.add_run(f"{prefix}{field}")
            _set_run_font(field_run, size=BODY_SIZE)

        details = item.get("details", [])
        if not isinstance(details, list):
            details = [details] if str(details).strip() else []
        if not grade:
            for detail in details:
                detail_text = str(detail).strip()
                if detail_text.lower().startswith("gpa"):
                    grade = detail_text
                    break

        if grade:
            line3 = document.add_paragraph()
            _style_paragraph(line3, after=2, line_spacing=LINE_SPACING)
            grade_value = grade if grade.lower().startswith("gpa") else f"GPA: {grade}"
            grade_run = line3.add_run(grade_value)
            grade_run.font.size = BODY_SIZE


def export_to_docx(result: dict, filename: str) -> None:
    document = Document()
    _configure_page(document)
    _set_global_font(document)

    resume = result["optimized_resume"]
    _add_two_column_header(document, resume.get("header", {}))
    _add_header_gap(document)

    sections = _section_map({"sections": sorted(resume.get("sections", []), key=_section_sort_key)})
    _add_summary(document, sections.get("summary"))
    _render_experience(document, sections.get("experience"))
    _render_skills(document, sections.get("skills"))
    _render_education(document, sections.get("education"))

    document.save(filename)
